from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "report_ieee.docx"


ACCENT = "087A74"
LIGHT = "EAF3F2"
LINE = "C9D6D4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            if i == 0 and header:
                set_cell_shading(cell, LIGHT)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(table)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(8, 122, 116) if level == 1 else RGBColor(21, 32, 33)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Dynamic Trend and Event Detection in News Streams")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(8, 122, 116)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("A Hybrid ML-Deep Learning Pipeline with Ablation and Dashboard Evaluation")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.italic = True

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Abhay Pratap Rana | Saloni Sharma | Vaishnav Verma\n").bold = True
    authors.add_run("School of Computer Science & AI, Rishihood University")

    add_heading(doc, "Abstract")
    add_para(
        doc,
        "This report presents a reproducible Python pipeline for dynamic trend and event detection on timestamped news. "
        "The system combines lexical drift from top terms, probabilistic topic drift from LDA, and semantic drift from "
        "MiniLM sentence embeddings. A hybrid fusion layer normalizes and combines these signals for adjacent time-bin "
        "transitions. The project includes same-split ablation comparing ML-only, DL-only, and Hybrid variants, plus a "
        "local dashboard for inspecting events, driver terms, topics, hybrid transition scores, and evaluation metrics.",
    )

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "News streams evolve continuously. Analysts need systems that highlight dominant words, latent themes, and meaningful "
        "shifts between time windows. Keyword counting is transparent but misses semantic changes; topic models are interpretable "
        "but bag-of-words based; neural sentence embeddings capture meaning but are less directly explainable. This project combines "
        "all three views into one hybrid event detector.",
    )
    add_bullets(
        doc,
        [
            "Open modular codebase runnable through src.main.",
            "Hybrid architecture combining lexical Jaccard drift, LDA Jensen-Shannon drift, and MiniLM semantic centroid drift.",
            "Same-split ablation study for ML-only, DL-only, and Hybrid variants.",
            "Interactive Flask dashboard for event, topic, transition, and ablation inspection.",
        ],
    )

    add_heading(doc, "2. Related Work")
    add_para(
        doc,
        "The project builds on LDA topic modeling, temporal topic analysis, sentence-transformer text representations, and topic "
        "detection and tracking research. LDA provides interpretable topic mixtures, while sentence transformers provide dense "
        "semantic representations useful for clustering and temporal semantic drift.",
    )

    add_heading(doc, "3. System Architecture and Data")
    add_para(
        doc,
        "The system loads a news CSV, preprocesses text, assigns documents to chronological time bins, then computes three aligned "
        "drift signals over adjacent time-bin transitions. These signals feed a weighted hybrid score used for event ranking, "
        "dashboard visualization, and ablation.",
    )
    add_table(
        doc,
        ["Stage", "Component", "Output"],
        [
            ["Input", "News CSV with timestamp, text, and category", "Document stream"],
            ["Preprocessing", "Tokenization, stopword removal, time binning", "Clean tokens and time_bin"],
            ["Lexical ML", "Top-k terms and Jaccard drift", "Lexical transition drift"],
            ["Probabilistic ML", "LDA topic mixtures and JSD", "Topic transition drift"],
            ["Neural DL", "MiniLM embeddings and centroid cosine distance", "Semantic transition drift"],
            ["Fusion", "0.30 lexical + 0.56 LDA + 0.14 semantic", "Hybrid event score"],
        ],
    )
    add_para(
        doc,
        "Dataset: Kaggle News Category Dataset containing HuffPost news items with category labels, publication dates, headlines, "
        "and short descriptions. The default sample uses Jan-Jun 2018, with Jan-Apr as training and May-Jun as testing.",
    )
    add_table(
        doc,
        ["Statistic", "Value"],
        [
            ["Documents", "8,749"],
            ["Approximate vocabulary", "18,858 tokens"],
            ["Total tokens", "134,268"],
            ["Mean tokens per document", "About 15.3"],
            ["Default time bin width", "7 days"],
        ],
    )

    add_heading(doc, "4. Methods and Metrics")
    add_heading(doc, "4.1 Lexical ML Branch", level=2)
    add_para(
        doc,
        "For each time bin, the system aggregates tokens and extracts top-k terms. Lexical drift is computed as 1 - Jaccard similarity "
        "between adjacent top-term sets.",
    )
    add_heading(doc, "4.2 Probabilistic ML Branch: LDA", level=2)
    add_para(
        doc,
        "LDA is trained on the training split. For every document, the model produces a topic proportion vector. Adjacent topic drift "
        "is measured using Jensen-Shannon distance between mean topic mixtures for neighboring bins.",
    )
    add_heading(doc, "4.3 Neural DL Branch: Sentence Embeddings", level=2)
    add_para(
        doc,
        "Texts are encoded with frozen sentence-transformers/all-MiniLM-L6-v2. The system computes one mean embedding centroid per "
        "time bin and measures semantic drift as 1 - cosine similarity between adjacent centroids.",
    )
    add_heading(doc, "4.4 Hybrid Fusion", level=2)
    add_para(
        doc,
        "All three drift signals are aligned on the same held-out transitions and min-max normalized. The implemented hybrid score is: "
        "H(t,t+1) = 0.30 * lexical + 0.56 * LDA + 0.14 * semantic.",
    )
    add_heading(doc, "4.5 Event Detection and Ablation", level=2)
    add_para(
        doc,
        "Event candidates are high-percentile drift spikes. The ablation study compares ML-only, DL-only, and Hybrid scores on the "
        "same held-out transitions using category-distribution drift as a weak validation target.",
    )

    add_heading(doc, "5. Experimental Results")
    add_table(
        doc,
        ["Model", "Precision", "Recall", "F1", "Spearman"],
        [
            ["ML-only", "reported by CLI/dashboard", "reported", "reported", "reported"],
            ["DL-only", "reported by CLI/dashboard", "reported", "reported", "reported"],
            ["Hybrid", "reported by CLI/dashboard", "reported", "reported", "reported"],
        ],
    )
    add_para(
        doc,
        "The dashboard and CLI produce the concrete values for the selected time window, percentile threshold, topic count, and "
        "embedding model. The key evaluation design is that all rows are measured on identical held-out transitions.",
    )
    add_table(
        doc,
        ["Example Metric", "Representative Value"],
        [
            ["Mean adjacent Jaccard", "About 0.351"],
            ["Mean adjacent LDA JSD", "About 0.048"],
            ["LDA log-perplexity", "About -9.06"],
            ["LDA u-mass coherence", "About -7.16"],
        ],
    )

    add_heading(doc, "6. Dashboard")
    add_para(
        doc,
        "The local Flask dashboard runs with python -m src.webapp and opens at http://127.0.0.1:8000. It includes controls for "
        "CSV path, time window, train/test split, LDA topics, event percentile, ablation percentile, stopwords, and embedding model. "
        "It visualizes KPIs, transition drift, ablation metrics, top hybrid transitions, latest terms, LDA topics, and detected events.",
    )

    add_heading(doc, "7. Discussion")
    add_para(
        doc,
        "The system is stronger than a disconnected ensemble because each branch addresses a specific weakness: lexical drift is "
        "interpretable, LDA topic drift captures topic-mixture movement, and MiniLM semantic drift captures meaning when vocabulary "
        "changes. The same-split ablation makes the comparison fair.",
    )
    add_para(
        doc,
        "Limitations include weak supervision through editor category drift, hand-specified fusion weights, fixed time bins, and a "
        "frozen neural encoder that may miss domain-specific event semantics.",
    )

    add_heading(doc, "8. Future Work")
    add_bullets(
        doc,
        [
            "Learn fusion weights from labeled event data.",
            "Add deeper diagnostic ablations by category, sequence length, and transition type.",
            "Add Docker or a setup script for turn-key reproducibility.",
            "Add CI/CD and optionally deploy the dashboard.",
            "Scale evaluation to the full news corpus.",
        ],
    )

    add_heading(doc, "9. Conclusion")
    add_para(
        doc,
        "This project implements a reproducible hybrid trend and event detector for timestamped news streams. It combines lexical, "
        "probabilistic, and neural drift signals into a hybrid score, supports same-split ablation, and includes an interactive UI "
        "for evaluation and demonstration.",
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "References")
    add_bullets(
        doc,
        [
            "Blei, Ng, and Jordan, Latent Dirichlet Allocation, 2003.",
            "Mimno et al., Optimizing semantic coherence in topic models, 2011.",
            "Reimers and Gurevych, Sentence-BERT, 2019.",
            "Wang and McCallum, Topics over Time, 2006.",
            "Misra, News Category Dataset, Kaggle, 2018.",
            "Allan, Topic Detection and Tracking, 2002.",
            "Pedregosa et al., Scikit-learn, 2011.",
            "Rehurek and Sojka, Gensim, 2010.",
        ],
    )

    header = doc.sections[0].header.paragraphs[0]
    header.text = "Dynamic Trend and Event Detector"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(97, 113, 115)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = "Hybrid ML + DL Report"
        footer.runs[0].font.name = "Arial"
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(97, 113, 115)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
